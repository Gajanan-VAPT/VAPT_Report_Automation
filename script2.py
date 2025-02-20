from docx import Document    #used to create and manipulate the word document 
from docx.shared import Pt   #used to set the font size of the text 
from docx.enum.text import WD_ALIGN_PARAGRAPH   # used to align the paragraph 
from docx.enum.section import WD_SECTION   # used to create a new page or section in the document 
from docx.oxml.shared import OxmlElement, qn  # used to handle the XML element such as cell colour
import csv  # use to read and parse the CSV 


def make_cell_text_bold(cell):
    """Helper function to make cell text bold"""
    paragraph = cell.paragraphs[0]
    run = paragraph.runs
    if not run:  # If there's no run, create one
        run = paragraph.add_run(cell.text)
        run.bold = True
        cell.text = ''  # Clear the cell text since we're using the run
    else:
        run[0].bold = True

def create_table(doc,heading):
    doc.add_heading(heading, level=2)

    # Row 0-1
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    table.cell(0, 0).width = Pt(150)
    table.cell(1, 0).width = Pt(150)
    table.cell(0, 1).width = Pt(350)
    table.cell(1, 1).width = Pt(350)
    
    # Set headers and make bold
    cells = [
        (0, 0, 'Finding ID'),
        (0, 1, 'Description')
    ]
    
    for row, col, text in cells:
        cell = table.cell(row, col)
        cell.text = text
        make_cell_text_bold(cell)
    
    table.cell(1, 0).text = ''
    table.cell(1, 1).text = ''

    # Apply shading
    for col in range(2):
        table_header = table.cell(0, col)
        tcPr = table_header._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), 'C0D4EC')
        tcShading.set(qn('w:themeTint'), '40')
        tcPr.append(tcShading)

    # Row 2-3
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False

    # Set cell widths
    for col in range(3):
        table.cell(0, col).width = Pt(150)
        table.cell(1, col).width = Pt(150)

    # Set headers and make bold
    cells = [
        (0, 0, 'CVS Score'),
        (0, 1, 'Risk Rating'),
        (0, 2, 'Remote Exploitability')
    ]
    
    for row, col, text in cells:
        cell = table.cell(row, col)
        cell.text = text
        make_cell_text_bold(cell)

    # Clear data cells
    for col in range(3):
        table.cell(1, col).text = ''

    # Apply shading
    for col in range(3):
        table_header = table.cell(0, col)
        tcPr = table_header._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), 'C0D4EC')
        tcShading.set(qn('w:themeTint'), '40')
        tcPr.append(tcShading)

    # Row 4-5
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    # Set cell widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(150)

    # Set headers and make bold
    cells = [
        (0, 0, 'Affected Resource'),
        (0, 1, 'Module Name')
    ]
    
    for row, col, text in cells:
        cell = table.cell(row, col)
        cell.text = text
        make_cell_text_bold(cell)

    # Clear data cells
    table.cell(1, 0).text = ''
    table.cell(1, 1).text = ''

    # Apply shading
    for col in range(2):
        table_header = table.cell(0, col)
        tcPr = table_header._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), 'C0D4EC')
        tcShading.set(qn('w:themeTint'), '40')
        tcPr.append(tcShading)

    # Row 6-7
    table = doc.add_table(rows=2, cols=1)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    # Set cell widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(300)

    # Set header and make bold
    cell = table.cell(0, 0)
    cell.text = 'Security Risk'
    make_cell_text_bold(cell)
    
    # Clear data cell
    table.cell(1, 0).text = ''

    # Apply shading
    table_header = table.cell(0, 0)
    tcPr = table_header._tc.get_or_add_tcPr()
    tcShading = OxmlElement('w:shd')
    tcShading.set(qn('w:fill'), 'C0D4EC')
    tcShading.set(qn('w:themeTint'), '40')
    tcPr.append(tcShading)

    # Row 8
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    # Set cell widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(150)

    # Set header and make bold
    cell = table.cell(0, 0)
    cell.text = 'Business Impact'
    make_cell_text_bold(cell)
    
    # Clear data cell
    table.cell(0, 1).text = ''

    # Row 9-10
    table = doc.add_table(rows=2, cols=1)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    # Set cell widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(300)

    # Set header and make bold
    cell = table.cell(0, 0)
    cell.text = 'Workaround / Mitigation'
    make_cell_text_bold(cell)
    
    # Clear data cell
    table.cell(1, 0).text = ''

    # Apply shading
    table_header = table.cell(0, 0)
    tcPr = table_header._tc.get_or_add_tcPr()
    tcShading = OxmlElement('w:shd')
    tcShading.set(qn('w:fill'), 'C0D4EC')
    tcShading.set(qn('w:themeTint'), '40')
    tcPr.append(tcShading)

    # Row 11-12
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    # Set cell widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(150)

    # Set headers and make bold
    cells = [
        (0, 0, 'Tool used'),
        (0, 1, 'References')
    ]
    
    for row, col, text in cells:
        cell = table.cell(row, col)
        cell.text = text
        make_cell_text_bold(cell)

    # Clear data cells
    table.cell(1, 0).text = ''
    table.cell(1, 1).text = ''

    # Apply shading
    for col in range(2):
        table_header = table.cell(0, col)
        tcPr = table_header._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), 'C0D4EC')
        tcShading.set(qn('w:themeTint'), '40')
        tcPr.append(tcShading)

    # Row 13-14
    table = doc.add_table(rows=2, cols=1)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    # Set cell widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(300)

    # Set header and make bold
    cell = table.cell(0, 0)
    cell.text = 'Proof of Concept (POC)'
    make_cell_text_bold(cell)
    
    # Clear data cell
    table.cell(1, 0).text = ''

    # Apply shading
    table_header = table.cell(0, 0)
    tcPr = table_header._tc.get_or_add_tcPr()
    tcShading = OxmlElement('w:shd')
    tcShading.set(qn('w:fill'), 'C0D4EC')
    tcShading.set(qn('w:themeTint'), '40')
    tcPr.append(tcShading)

KEYWORDS ={
    'Apache': 'Apache',
    'Window': 'Windows',
    'SSH': 'SSH',
    'Oracle': 'Oracle',
}

def get_module_name(name, predefined_keywords):
    # Iterate through the predefined keywords and check if they match the name
    for keyword in predefined_keywords:
        if keyword.lower() in name.lower():
            return keyword
    return ""  # Return empty string if no match is found
def append_data(doc, csv_row_no, data_to_append, predefined_keywords):
    # table - 0
    # finding_id
    doc.tables[csv_row_no*8].rows[1].cells[0].text = data_to_append['finding_id']
    # Prepend the required text to the description
    description = f"During Vulnerability assessment and Penetration testing we observed that, {data_to_append['description']}"

# Then update the document with the new description
    doc.tables[csv_row_no*8].rows[1].cells[1].text = description


    # table - 1
    # cvs_score
    doc.tables[csv_row_no*8+1].rows[1].cells[0].text = data_to_append['cvs_score']
    # risk_rating
    cell = doc.tables[csv_row_no*8+1].rows[1].cells[1]
    cell.text = data_to_append['risk_factor']
    if data_to_append['risk_factor'].lower() == 'critical':
        tcPr = cell._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), '800000')
        tcPr.append(tcShading)
    elif data_to_append['risk_factor'].lower() == 'high':
        tcPr = cell._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), 'FF6600')
        tcPr.append(tcShading)
    elif data_to_append['risk_factor'].lower() == 'medium':
        tcPr = cell._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), 'FFFF00')
        tcPr.append(tcShading)
    elif data_to_append['risk_factor'].lower() == 'low':
        tcPr = cell._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), '008000')
        tcPr.append(tcShading)
    elif data_to_append['risk_factor'].lower() == 'informational':
        tcPr = cell._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), '3366FF')
        tcPr.append(tcShading)

    # remote_exploitability
    doc.tables[csv_row_no*8+1].rows[1].cells[2].text = "Yes"

    # table - 2
    # affected_resource
    doc.tables[csv_row_no*8+2].rows[1].cells[0].text = data_to_append['affected_resource']
    # module_name - dynamically matched from predefined keywords
    module_name = get_module_name(data_to_append['name'], predefined_keywords)
    doc.tables[csv_row_no*8+2].rows[1].cells[1].text = module_name

    # table - 3
    # security_risk
    doc.tables[csv_row_no*8+3].rows[1].cells[0].text = ""

    # table - 4w
    # business_impact
    doc.tables[csv_row_no*8+4].rows[0].cells[1].text = ""

    # table - 5
    # workaround_mitigation
    solution = f"It is recommended: \n-To {data_to_append['mitigation']}"
    doc.tables[csv_row_no*8+5].rows[1].cells[0].text = solution

    # table - 6
    # tool_used
    doc.tables[csv_row_no*8+6].rows[1].cells[0].text = "Nessus"
    # references
    doc.tables[csv_row_no*8+6].rows[1].cells[1].text = data_to_append['references']

    # table - 7 (Proof of Concept)
    proof_of_concept_cell = doc.tables[csv_row_no*8+7].rows[1].cells[0]
    existing_text = proof_of_concept_cell.text  # Preserve existing content
    proof_of_concept_cell.text = ""  # Clear the cell to format text properly

    # Adding "Figure 1 - Shows" with "Figure 1" in bold
    proof_paragraph = proof_of_concept_cell.paragraphs[0]
    run = proof_paragraph.add_run("\nFigure 1")
    run.bold = True
    proof_paragraph.add_run(" - Shows" + existing_text)  # Append previous content

def main():
    # Create a new document
    doc = Document()

    # Add a title
    title = doc.add_heading('RNS DATA AUTOMATION', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    section = doc.sections[-1]
    csv_file_path = 'dataset.csv'

    # Counter for finding IDs
    finding_id_counter = 1

    # Open the CSV file and read its contents
    with open(csv_file_path, 'r', encoding='utf-8', errors='ignore') as csv_file:
        csv_reader = csv.reader(csv_file)
        # Skip the header row if needed
        next(csv_reader)

        for csv_row_no, csv_row_data in enumerate(csv_reader):
            # Generate auto-incrementing finding ID
            finding_id = f"ABCXYZ-{finding_id_counter}"
            finding_id_counter += 1  # Increment counter

            # Prepare data to append with module name from CSV data
            data_to_append = {
                "name": csv_row_data[7],  # Vulnerability name
                "finding_id": finding_id,  # Auto-incrementing ID
                "description": csv_row_data[9].split('.')[0],
                "cvs_score": csv_row_data[2],
                "risk_rating": "",
                "risk_factor": csv_row_data[18],
                "remote_exploitability": "Yes",
                "affected_resource": [],
                "module_name": "",  # Placeholder, this will be set dynamically
                "security_risk": "",
                "business_impact": "",
                "mitigation": csv_row_data[10],
                "tool_used": "Nessus",
                "references": csv_row_data[11].split('\n')[:1],
                "proof_of_concept": "",
            }

            # Create table for each row
            create_table(doc, data_to_append['name'])
            doc.add_section(WD_SECTION.NEW_PAGE)
            append_data(doc, csv_row_no, data_to_append, KEYWORDS)

    # Save the document
    doc.save('output_document.docx')

if __name__ == "__main__":
    main()
