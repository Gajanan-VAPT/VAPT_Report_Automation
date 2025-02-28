from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import csv
import logging  
import os
import sys    #used to print the log to console
from datetime import datetime

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler(sys.stdout)]
)

logger = logging.getLogger(__name__)
# Set up logging
  #Creates a logger object  ( logger will be used in the entire script now)

def set_document_font(doc, font_name="Helvetica", font_size=11):
    """Set the default font for the entire document."""
    logger.info(f"Setting document font to {font_name}, size {font_size}")   # This will log and message in the console
    try:
        style = doc.styles['Normal']
        style.font.name = font_name        # Used try catch for exception handling
        style.font.size = Pt(font_size)
        
        # Ensure the correct font is applied to tables too
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
        logger.info("Document font set successfully")    # If this execute this will print font set successfully
    except Exception as e:
        logger.error(f"Failed to set document font: {str(e)}")   
        raise

def make_cell_text_bold(cell):
    """Helper function to make cell text bold"""
    try:
        paragraph = cell.paragraphs[0]
        run = paragraph.runs
        if not run:  # If there's no run, create one
            run = paragraph.add_run(cell.text)
            run.bold = True
            cell.text = ''  # Clear the cell text since we're using the run
        else:
            run[0].bold = True
    except Exception as e:
        logger.error(f"Failed to make cell text bold: {str(e)}")
        raise

def format_affected_resources(table, affected_resources):
    logger.info(f"Formatting affected resources table for {len(affected_resources)} resources")
    try:
        cell = table.cell(1, 0)
        cell.text = ""  # Clear existing content
        
        # Sort the resources for consistent display
        resources = sorted(affected_resources)
        total = len(resources)
        
        # Clear all existing paragraphs in the cell except the first one
        for p in cell.paragraphs[1:]:
            p._element.getparent().remove(p._element)
        
        # Get the first paragraph
        first_paragraph = cell.paragraphs[0]
        first_paragraph.text = ""  # Ensure it's empty
        
        # For fewer than 5 resources: vertical layout (single column)
        if total < 5:
            # Add first resource to the existing paragraph
            if resources:
                run = first_paragraph.add_run(resources[0])
                run.font.size = Pt(10.5)
            
            # Add remaining resources in separate paragraphs
            for i in range(1, total):
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(resources[i])
                run.font.size = Pt(10.5)
                
                # Set paragraph spacing
                paragraph.space_after = Pt(0)
                paragraph.space_before = Pt(0)
        
        # For 5 or more resources: two-column layout
        else:
            # Calculate number of rows needed (ceiling division)
            rows = (total + 1) // 2  # This ensures we round up
            
            # Add first row to existing paragraph
            if resources:
                run = first_paragraph.add_run(resources[0])
                run.font.size = Pt(10)
                padding = 35 - len(resources[0])  # Adjust 35 based on your needs
                run = first_paragraph.add_run(" " * padding)
                
                # Add right column for first row if it exists
                if rows < len(resources):
                    run = first_paragraph.add_run(resources[rows])
                    run.font.size = Pt(10)
            
            # Set paragraph spacing
            first_paragraph.space_after = Pt(0)
            first_paragraph.space_before = Pt(0)
            
            # Create paragraph for each remaining row (starting from index 1)
            for i in range(1, rows):
                paragraph = cell.add_paragraph()
                
                # Add left column
                if i < len(resources):
                    run = paragraph.add_run(resources[i])
                    run.font.size = Pt(10)  
                    padding = 35 - len(resources[i])  # Adjust based on your needs
                    run = paragraph.add_run(" " * padding)
                    
                    # Add right column if it exists
                    if i + rows < len(resources):
                        run = paragraph.add_run(resources[i + rows])
                        run.font.size = Pt(10)
                
                # Set paragraph spacing
                paragraph.space_after = Pt(0)
                paragraph.space_before = Pt(0)
        logger.info("Successfully formatted affected resources")
    except Exception as e:
        logger.error(f"Failed to format affected resources: {str(e)}")
        raise

def create_table(doc, heading):
    logger.info(f"Creating table with heading: {heading}")
    try:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(heading)
        run.bold = True
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(0)
        # Apply the theme color (matches Word's default heading color)
        rPr = run._element.get_or_add_rPr()
        color = parse_xml(r'<w:color {} w:val="365F91"/>'.format(nsdecls('w')))  # Blue theme color
        rPr.append(color)

        # Row 0-1
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        table.cell(0, 0).width = Pt(150)
        table.cell(1, 0).width = Pt(150)
        table.cell(0, 1).width = Pt(350)
        table.cell(1, 1).width = Pt(350)
        
        # Set headers and make bold
        cells = [
            (0, 0, 'Finding ID'),
            (0, 1, 'Description')
        ]
        
        for row, col, text in cells:
            cell = table.cell(row, col)
            cell.text = text
            make_cell_text_bold(cell)
        
        table.cell(1, 0).text = ''
        table.cell(1, 1).text = ''

        # Apply shading
        for col in range(2):
            table_header = table.cell(0, col)
            tcPr = table_header._tc.get_or_add_tcPr()
            tcShading = OxmlElement('w:shd')
            tcShading.set(qn('w:fill'), 'C0D4EC')
            tcShading.set(qn('w:themeTint'), '40')
            tcPr.append(tcShading)

        # Row 2-3
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False

        # Set cell widths
        for col in range(3):
            table.cell(0, col).width = Pt(150)
            table.cell(1, col).width = Pt(150)

        # Set headers and make bold
        cells = [
            (0, 0, 'CVS Score'),
            (0, 1, 'Risk Rating'),
            (0, 2, 'Remote Exploitability')
        ]
        
        for row, col, text in cells:
            cell = table.cell(row, col)
            cell.text = text
            make_cell_text_bold(cell)

        # Clear data cells
        for col in range(3):
            table.cell(1, col).text = ''

        # Apply shading
        for col in range(3):
            table_header = table.cell(0, col)
            tcPr = table_header._tc.get_or_add_tcPr()
            tcShading = OxmlElement('w:shd')
            tcShading.set(qn('w:fill'), 'C0D4EC')
            tcShading.set(qn('w:themeTint'), '40')
            tcPr.append(tcShading)

        # Row 4-5
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        
        # Set cell widths
        for row in table.rows:
            for cell in row.cells:
                cell.width = Pt(150)
        # Merge column 0 and column 1 to create space for "Affected Resource"
        table.cell(0, 0).merge(table.cell(0, 1))  # Merge first row, columns 0 and 1
        table.cell(1, 0).merge(table.cell(1, 1))  # Merge second row, columns 0 and 1
        # Set headers and make bold
        cells = [
            (0, 0, 'Affected Resource'),
            (0, 2, 'Module Name')
        ]
        
        for row, col, text in cells:
            cell = table.cell(row, col)
            cell.text = text
            make_cell_text_bold(cell)

        # Clear data cells
        table.cell(1, 0).text = ''
        table.cell(1, 1).text = ''

        # Apply shading
        for col in range(3):
            table_header = table.cell(0, col)
            tcPr = table_header._tc.get_or_add_tcPr()
            tcShading = OxmlElement('w:shd')
            tcShading.set(qn('w:fill'), 'C0D4EC')
            tcShading.set(qn('w:themeTint'), '40')
            tcPr.append(tcShading)

        # Row 6-7
        table = doc.add_table(rows=2, cols=1)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        
        # Set cell widths
        for row in table.rows:
            for cell in row.cells:
                cell.width = Pt(300)

        # Set header and make bold
        cell = table.cell(0, 0)
        cell.text = 'Security Risk'
        make_cell_text_bold(cell)
        
        # Clear data cell
        table.cell(1, 0).text = ''

        # Apply shading
        table_header = table.cell(0, 0)
        tcPr = table_header._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), 'C0D4EC')
        tcShading.set(qn('w:themeTint'), '40')
        tcPr.append(tcShading)

        # Row 8
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        
        # Set cell widths
        for row in table.rows:
            for cell in row.cells:
                cell.width = Pt(150)

        # Set header and make bold
        cell = table.cell(0, 0)
        cell.text = 'Business Impact'
        make_cell_text_bold(cell)
        
        # Clear data cell
        table.cell(0, 1).text = ''

        # Row 9-10
        table = doc.add_table(rows=2, cols=1)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        
        # Set cell widths
        for row in table.rows:
            for cell in row.cells:
                cell.width = Pt(300)

        # Set header and make bold
        cell = table.cell(0, 0)
        cell.text = 'Workaround / Mitigation'
        make_cell_text_bold(cell)
        
        # Clear data cell
        table.cell(1, 0).text = ''

        # Apply shading
        table_header = table.cell(0, 0)
        tcPr = table_header._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), 'C0D4EC')
        tcShading.set(qn('w:themeTint'), '40')
        tcPr.append(tcShading)

        # Row 11-12
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        
        # Set cell widths
        for row in table.rows:
            for cell in row.cells:
                cell.width = Pt(150)
        table.cell(0,0).merge(table.cell(0,1))
        table.cell(1, 0).merge(table.cell(1,1)) 
        # Set headers and make bold
        cells = [
            (0, 0, 'Tool used'),
            (0, 2, 'References')
        ]
        
        for row, col, text in cells:
            cell = table.cell(row, col)
            cell.text = text
            make_cell_text_bold(cell)

        # Clear data cells
        table.cell(1, 0).text = ''
        table.cell(1, 1).text = ''

        # Apply shading
        for col in range(3):
            table_header = table.cell(0, col)
            tcPr = table_header._tc.get_or_add_tcPr()
            tcShading = OxmlElement('w:shd')
            tcShading.set(qn('w:fill'), 'C0D4EC')
            tcShading.set(qn('w:themeTint'), '40')
            tcPr.append(tcShading)

        # Row 13-14
        table = doc.add_table(rows=2, cols=1)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        
        # Set cell widths
        for row in table.rows:
            for cell in row.cells:
                cell.width = Pt(300)

        # Set header and make bold
        cell = table.cell(0, 0)
        cell.text = 'Proof of Concept (POC)'
        make_cell_text_bold(cell)
        
        # Clear data cell
        table.cell(1, 0).text = ''

        # Apply shading
        table_header = table.cell(0, 0)
        tcPr = table_header._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), 'C0D4EC')
        tcShading.set(qn('w:themeTint'), '40')
        tcPr.append(tcShading)
        
        logger.info(f"Table created successfully for heading: {heading}")
        return True
    except Exception as e:
        logger.error(f"Failed to create table for {heading}: {str(e)}")
        raise

KEYWORDS = {
    'Apache': 'Apache',
    'Window': 'Windows',
    'SSH': 'SSH',
    'Oracle': 'Oracle',
    'DNS': 'DNS',
}

def get_module_name(name, predefined_keywords):
    """Helper function to match module names from predefined keywords"""
    logger.debug(f"Finding module name for: {name}")
    try:
        # Iterate through the predefined keywords and check if they match the name
        for keyword in predefined_keywords:
            if keyword.lower() in name.lower():
                logger.debug(f"Found matching module name: {keyword}")
                return keyword
        logger.debug("No matching module name found")
        return ""  # Return empty string if no match is found
    except Exception as e:
        logger.error(f"Error in get_module_name: {str(e)}")
        return ""

def append_data(doc, csv_row_no, data_to_append, predefined_keywords):
    logger.info(f"Appending data for vulnerability {data_to_append['finding_id']}")
    try:
        # table - 0
        # finding_id
        finding_id_cell = doc.tables[csv_row_no * 8].rows[1].cells[0]
        finding_id_cell.text = data_to_append['finding_id']
        finding_id_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Prepend the required text to the description
        description = f"During Vulnerability assessment and Penetration testing we observed that, {data_to_append['description']}"
        doc.tables[csv_row_no * 8].rows[1].cells[1].text = description

        # table - 1
        # cvs_score
        doc.tables[csv_row_no * 8 + 1].rows[1].cells[0].text = data_to_append['cvs_score']

        # risk_rating (with color highlighting)
        cell = doc.tables[csv_row_no * 8 + 1].rows[1].cells[1]
        cell.text = data_to_append['risk_factor'].capitalize()

        risk_colors = {
            "critical": "800000",
            "high": "FFC404",
            "medium": "FFFF00",
            "low": "008000",
            "informational": "3366FF",
        }
        
        risk_color = risk_colors.get(data_to_append['risk_factor'].lower())
        if risk_color:
            tcPr = cell._tc.get_or_add_tcPr()
            tcShading = OxmlElement('w:shd')
            tcShading.set(qn('w:fill'), risk_color)
            tcPr.append(tcShading)

        # remote_exploitability
        doc.tables[csv_row_no * 8 + 1].rows[1].cells[2].text = "Yes"

        # table - 2 (Affected Resource & Module Name)
        table = doc.tables[csv_row_no * 8 + 2]
        
        # Get affected resources as a list
        affected_resources = data_to_append["affected_resource"].split("\n")
        
        # Use new formatting function
        format_affected_resources(table, affected_resources)

        # Module Name - dynamically matched from predefined keywords
        module_name_cell = table.cell(1, 2)
        module_name_cell.text = get_module_name(data_to_append['name'], predefined_keywords)
        module_name_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # table - 3 (Security Risk)
        doc.tables[csv_row_no * 8 + 3].rows[1].cells[0].text = ""

        # table - 4 (Business Impact)
        doc.tables[csv_row_no * 8 + 4].rows[0].cells[1].text = ""

        # table - 5 (Workaround / Mitigation)
        solution = f"It is recommended: \n-To {data_to_append['mitigation']}"
        doc.tables[csv_row_no * 8 + 5].rows[1].cells[0].text = solution

        # table - 6 (Tool Used & References)
        doc.tables[csv_row_no * 8 + 6].rows[1].cells[0].text = "Nessus"
        doc.tables[csv_row_no * 8 + 6].rows[1].cells[2].text = data_to_append['references']

        # table - 7 (Proof of Concept)
        proof_of_concept_cell = doc.tables[csv_row_no * 8 + 7].rows[1].cells[0]
        existing_text = proof_of_concept_cell.text  # Preserve existing content
        proof_of_concept_cell.text = ""  # Clear the cell to format text properly

        # Adding "Figure 1 - Shows" with "Figure 1" in bold
        proof_paragraph = proof_of_concept_cell.paragraphs[0]
        run = proof_paragraph.add_run("\nFigure 1")
        run.bold = True
        proof_paragraph.add_run(" - Shows " + existing_text)
        
        logger.info(f"Successfully appended data for vulnerability {data_to_append['finding_id']}")
        return True
    except Exception as e:
        logger.error(f"Failed to append data for {data_to_append['finding_id']}: {str(e)}")
        raise

def validate_csv_columns(column_map, required_fields):
    """Validate if all required columns exist in the CSV file"""
    logger.info("Validating CSV columns")
    missing_fields = []
    for field, column_name in required_fields.items():
        if column_name not in column_map:
            missing_fields.append(column_name)
    
    if missing_fields:
        error_msg = f"Missing required columns in CSV: {', '.join(missing_fields)}"
        logger.error(error_msg)
        return False, error_msg
    
    logger.info("All required columns found in CSV")
    return True, ""

def main():
    logger.info("Starting document creation process")
    
    try:
        # Create a new document
        doc = Document()
        logger.info("Created new document")
        
        set_document_font(doc, "Helvetica", 10.5)
        grouped_vulnerabilities = {}  

        # Add a title
        title = doc.add_heading('RNS DATA AUTOMATION', level=1)
        title.alignment = 1  # Center align
        logger.info("Added document title")
        
        csv_file_path = 'dataset.csv'
        logger.info(f"Using CSV file: {csv_file_path}")

        # Check if CSV file exists
        if not os.path.exists(csv_file_path):
            logger.error(f"CSV file not found: {csv_file_path}")
            raise FileNotFoundError(f"CSV file not found: {csv_file_path}")

        # Counter for finding IDs
        finding_id_counter = 1

        # Define the required fields and their corresponding column names
        required_fields = {
            'name': 'Name',
            'description': 'Description',
            'cvs_score': 'CVSS v3.0 Base Score',
            'risk_factor': 'Risk Factor',
            'host': 'Host',
            'port': 'Port',
            'mitigation': 'Solution',
            'references': 'See Also'
        }

        # Open the CSV file and read its contents
        try:
            with open(csv_file_path, 'r', encoding='utf-8', errors='ignore') as csv_file:
                logger.info("Successfully opened CSV file")
                
                # Use a try-except block for CSV reading operations
                try:
                    csv_reader = csv.reader(csv_file)
                    
                    # Read the header row to determine column indices
                    header = next(csv_reader)
                    logger.info(f"CSV header read successfully with {len(header)} columns")
                    
                    # Create a mapping of column names to indices
                    column_map = {}
                    for i, column_name in enumerate(header):
                        column_map[column_name] = i
                    
                    # Validate CSV columns
                    valid, error_message = validate_csv_columns(column_map, required_fields)
                    if not valid:
                        logger.error(f"CSV validation failed: {error_message}")
                        logger.info(f"Available columns: {', '.join(column_map.keys())}")
                        raise ValueError(error_message)
                    
                    # Process each row in the CSV
                    row_count = 0
                    for csv_row_data in csv_reader:
                        row_count += 1
                        logger.info(f"Processing row {row_count}")
                        
                        try:
                            # Extract data using column names instead of hardcoded indices
                            vulnerability_name = csv_row_data[column_map[required_fields['name']]]
                            host = csv_row_data[column_map[required_fields['host']]]
                            port = csv_row_data[column_map[required_fields['port']]]
                            affected_host = f"{host}:{port}"
                            
                            # Get other required fields
                            description = csv_row_data[column_map[required_fields['description']]].split('.')[0]
                            cvs_score = csv_row_data[column_map[required_fields['cvs_score']]]
                            risk_factor = csv_row_data[column_map[required_fields['risk_factor']]]
                            mitigation = csv_row_data[column_map[required_fields['mitigation']]]
                            
                            # Handle possible empty references
                            references_idx = column_map[required_fields['references']]
                            references = csv_row_data[references_idx].split('\n')[0] if csv_row_data[references_idx] else ""
                        
                            if vulnerability_name in grouped_vulnerabilities:
                                grouped_vulnerabilities[vulnerability_name]["affected_resource"].add(affected_host)
                            else:
                                # Assign a new finding ID and store all required details
                                grouped_vulnerabilities[vulnerability_name] = {
                                    "name": vulnerability_name,
                                    "finding_id": f"ABCXYZ-{finding_id_counter}",
                                    "description": description,
                                    "cvs_score": cvs_score,
                                    "risk_factor": risk_factor,
                                    "remote_exploitability": "Yes",
                                    "affected_resource": {affected_host},
                                    "mitigation": mitigation,
                                    "references": references,
                                }
                                finding_id_counter += 1  # Increment counter
                            
                            logger.info(f"Successfully processed vulnerability: {vulnerability_name}")
                        
                        except IndexError as e:
                            logger.error(f"Invalid data in row {row_count}: {str(e)}")
                            logger.warning(f"Skipping row {row_count} due to missing or invalid data")
                            continue
                        except Exception as e:
                            logger.error(f"Error processing row {row_count}: {str(e)}")
                            logger.warning(f"Skipping row {row_count} due to processing error")
                            continue 
                
                except csv.Error as e:
                    logger.error(f"CSV parsing error: {str(e)}")
                    raise
        
        except FileNotFoundError:
            logger.error(f"CSV file not found: {csv_file_path}")
            raise
        except PermissionError:
            logger.error(f"Permission denied when accessing CSV file: {csv_file_path}")
            raise
        except Exception as e:
            logger.error(f"Unexpected error when reading CSV file: {str(e)}")
            raise

        # Check if we have any vulnerabilities to process
        if not grouped_vulnerabilities:
            logger.warning("No valid vulnerabilities found in the CSV file")
            raise ValueError("No valid vulnerabilities found in the CSV file")

        logger.info(f"Found {len(grouped_vulnerabilities)} unique vulnerabilities")

        # Create tables for each vulnerability
        for index, (vulnerability_name, data_to_append) in enumerate(grouped_vulnerabilities.items()):
            logger.info(f"Creating table for vulnerability {index + 1}: {vulnerability_name}")
            
            # Fix: Strip spaces before joining affected resources
            data_to_append["affected_resource"] = "\n".join(res.strip() for res in sorted(data_to_append["affected_resource"]))

            # Create the table
            numbered_finding_name = f"{index + 1}. {vulnerability_name}"
            create_table(doc, numbered_finding_name)
        
            # Add new page except for the last vulnerability
            if index < len(grouped_vulnerabilities) - 1:
                doc.add_section(WD_SECTION.NEW_PAGE)
                logger.info("Added new page for next vulnerability")
            
            append_data(doc, index, data_to_append, KEYWORDS)

        # Save the document
        try:
            doc.save('output_document.docx')
            logger.info("Document saved successfully as 'output_document.docx'")
        except PermissionError:
            logger.error("Permission denied when saving the document. Check if the file is open in another application.")
            raise
        except Exception as e:
            logger.error(f"Failed to save document: {str(e)}")
            raise

        logger.info("Document creation process completed successfully")
        return True

    except Exception as e:
        logger.critical(f"Document creation failed: {str(e)}")
        return False

if __name__ == "__main__":
    try:
        success = main()
        if success:
            print("Document created successfully!")
            logger.info("Script execution completed successfully")
        else:
            print("Document creation failed. Check the log file for details.")
            logger.error("Script execution failed")
    except Exception as e:
        print(f"An unexpected error occurred: {str(e)}")
        logger.critical(f"Unhandled exception: {str(e)}", exc_info=True)