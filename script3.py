from docx import Document    #used to create and manipulate the word document 
from docx.shared import Pt   #used to set the font size of the text 
from docx.enum.text import WD_ALIGN_PARAGRAPH   # used to align the paragraph 
from docx.enum.section import WD_SECTION   # used to create a new page or section in the document 
from docx.oxml.shared import OxmlElement, qn  # used to handle the XML element such as cell colour
from docx.enum.table import WD_ALIGN_VERTICAL 
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import csv  # use to read and parse the CSV 

def set_document_font(doc, font_name="Helvetica", font_size=11):
    """Set the default font for the entire document."""
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)  # Set default size (optional)
    
    # Ensure the correct font is applied to tables too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
def make_cell_text_bold(cell):
    """Helper function to make cell text bold"""
    paragraph = cell.paragraphs[0]
    run = paragraph.runs
    if not run:  # If there's no run, create one
        run = paragraph.add_run(cell.text)
        run.bold = True
        cell.text = ''  # Clear the cell text since we're using the run
    else:
        run[0].bold = True

def format_affected_resources(table, affected_resources):
    """
    Format affected resources with conditional layout:
    - If fewer than 5 resources: display in a single vertical column
    - If 5 or more resources: display in a two-column layout
    
    Args:
        table: The document table to modify
        affected_resources: List of affected resource strings
    """
    cell = table.cell(1, 0)
    cell.text = ""  # Clear existing content
    
    # Sort the resources for consistent display
    resources = sorted(affected_resources)
    total = len(resources)
    
    # Clear all existing paragraphs in the cell except the first one
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    
    # Get the first paragraph
    first_paragraph = cell.paragraphs[0]
    first_paragraph.text = ""  # Ensure it's empty
    
    # For fewer than 5 resources: vertical layout (single column)
    if total <= 5:
        # Add first resource to the existing paragraph
        if resources:
            run = first_paragraph.add_run(resources[0])
            run.font.size = Pt(10)
        
        # Add remaining resources in separate paragraphs
        for i in range(1, total):
            paragraph = cell.add_paragraph()
            run = paragraph.add_run(resources[i])
            run.font.size = Pt(10)
            
            # Set paragraph spacing
            paragraph.space_after = Pt(0)
            paragraph.space_before = Pt(0)
    
    # For 5 or more resources: two-column layout
    else:
        # Calculate number of rows needed (ceiling division)
        rows = (total + 1) // 2  # This ensures we round up
        
        # Add first row to existing paragraph
        if resources:
            run = first_paragraph.add_run(resources[0])
            run.font.size = Pt(10)
            padding = 35 - len(resources[0])  # Adjust 35 based on your needs
            run = first_paragraph.add_run(" " * padding)
            
            # Add right column for first row if it exists
            if rows < len(resources):
                run = first_paragraph.add_run(resources[rows])
                run.font.size = Pt(10)
        
        # Set paragraph spacing
        first_paragraph.space_after = Pt(0)
        first_paragraph.space_before = Pt(0)
        
        # Create paragraph for each remaining row (starting from index 1)
        for i in range(1, rows):
            paragraph = cell.add_paragraph()
            
            # Add left column
            if i < len(resources):
                run = paragraph.add_run(resources[i])
                run.font.size = Pt(10)  
                padding = 35 - len(resources[i])  # Adjust based on your needs
                run = paragraph.add_run(" " * padding)
                
                # Add right column if it exists
                if i + rows < len(resources):
                    run = paragraph.add_run(resources[i + rows])
                    run.font.size = Pt(10)
            
            # Set paragraph spacing
            paragraph.space_after = Pt(0)
            paragraph.space_before = Pt(0)

def create_table(doc, heading):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(heading)
    run.bold = True
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(0)
    # Apply the theme color (matches Word's default heading color)
    rPr = run._element.get_or_add_rPr()
    color = parse_xml(r'<w:color {} w:val="365F91"/>'.format(nsdecls('w')))  # Blue theme color
    rPr.append(color)

    # Row 0-1
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    table.cell(0, 0).width = Pt(150)
    table.cell(1, 0).width = Pt(150)
    table.cell(0, 1).width = Pt(350)
    table.cell(1, 1).width = Pt(350)
    
    # Set headers and make bold
    cells = [
        (0, 0, 'Finding ID'),
        (0, 1, 'Description')
    ]
    
    for row, col, text in cells:
        cell = table.cell(row, col)
        cell.text = text
        make_cell_text_bold(cell)
    
    table.cell(1, 0).text = ''
    table.cell(1, 1).text = ''

    # Apply shading
    for col in range(2):
        table_header = table.cell(0, col)
        tcPr = table_header._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), 'C0D4EC')
        tcShading.set(qn('w:themeTint'), '40')
        tcPr.append(tcShading)

    # Row 2-3
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False

    # Set cell widths
    for col in range(3):
        table.cell(0, col).width = Pt(150)
        table.cell(1, col).width = Pt(150)

    # Set headers and make bold
    cells = [
        (0, 0, 'CVS Score'),
        (0, 1, 'Risk Rating'),
        (0, 2, 'Remote Exploitability')
    ]
    
    for row, col, text in cells:
        cell = table.cell(row, col)
        cell.text = text
        make_cell_text_bold(cell)

    # Clear data cells
    for col in range(3):
        table.cell(1, col).text = ''

    # Apply shading
    for col in range(3):
        table_header = table.cell(0, col)
        tcPr = table_header._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), 'C0D4EC')
        tcShading.set(qn('w:themeTint'), '40')
        tcPr.append(tcShading)

    # Row 4-5
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    # Set cell widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(150)
    # Merge column 0 and column 1 to create space for "Affected Resource"
    table.cell(0, 0).merge(table.cell(0, 1))  # Merge first row, columns 0 and 1
    table.cell(1, 0).merge(table.cell(1, 1))  # Merge second row, columns 0 and 1
    # Set headers and make bold
    cells = [
        (0, 0, 'Affected Resource'),
        (0, 2, 'Module Name')
    ]
    
    for row, col, text in cells:
        cell = table.cell(row, col)
        cell.text = text
        make_cell_text_bold(cell)

    # Clear data cells
    table.cell(1, 0).text = ''
    table.cell(1, 1).text = ''

    # Apply shading
    for col in range(3):
        table_header = table.cell(0, col)
        tcPr = table_header._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), 'C0D4EC')
        tcShading.set(qn('w:themeTint'), '40')
        tcPr.append(tcShading)

    # Row 6-7
    table = doc.add_table(rows=2, cols=1)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    # Set cell widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(300)

    # Set header and make bold
    cell = table.cell(0, 0)
    cell.text = 'Security Risk'
    make_cell_text_bold(cell)
    
    # Clear data cell
    table.cell(1, 0).text = ''

    # Apply shading
    table_header = table.cell(0, 0)
    tcPr = table_header._tc.get_or_add_tcPr()
    tcShading = OxmlElement('w:shd')
    tcShading.set(qn('w:fill'), 'C0D4EC')
    tcShading.set(qn('w:themeTint'), '40')
    tcPr.append(tcShading)

    # Row 8
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    # Set cell widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(150)

    # Set header and make bold
    cell = table.cell(0, 0)
    cell.text = 'Business Impact'
    make_cell_text_bold(cell)
    
    # Clear data cell
    table.cell(0, 1).text = ''

    # Row 9-10
    table = doc.add_table(rows=2, cols=1)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    # Set cell widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(300)

    # Set header and make bold
    cell = table.cell(0, 0)
    cell.text = 'Workaround / Mitigation'
    make_cell_text_bold(cell)
    
    # Clear data cell
    table.cell(1, 0).text = ''

    # Apply shading
    table_header = table.cell(0, 0)
    tcPr = table_header._tc.get_or_add_tcPr()
    tcShading = OxmlElement('w:shd')
    tcShading.set(qn('w:fill'), 'C0D4EC')
    tcShading.set(qn('w:themeTint'), '40')
    tcPr.append(tcShading)

    # Row 11-12
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    # Set cell widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(150)
    table.cell(0,0).merge(table.cell(0,1))
    table.cell(1, 0).merge(table.cell(1,1)) 
    # Set headers and make bold
    cells = [
        (0, 0, 'Tool used'),
        (0, 2, 'References')
    ]
    
    for row, col, text in cells:
        cell = table.cell(row, col)
        cell.text = text
        make_cell_text_bold(cell)

    # Clear data cells
    table.cell(1, 0).text = ''
    table.cell(1, 1).text = ''

    # Apply shading
    for col in range(3):
        table_header = table.cell(0, col)
        tcPr = table_header._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), 'C0D4EC')
        tcShading.set(qn('w:themeTint'), '40')
        tcPr.append(tcShading)

    # Row 13-14
    table = doc.add_table(rows=2, cols=1)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    # Set cell widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(300)

    # Set header and make bold
    cell = table.cell(0, 0)
    cell.text = 'Proof of Concept (POC)'
    make_cell_text_bold(cell)
    
    # Clear data cell
    table.cell(1, 0).text = ''

    # Apply shading
    table_header = table.cell(0, 0)
    tcPr = table_header._tc.get_or_add_tcPr()
    tcShading = OxmlElement('w:shd')
    tcShading.set(qn('w:fill'), 'C0D4EC')
    tcShading.set(qn('w:themeTint'), '40')
    tcPr.append(tcShading)

KEYWORDS = {
    'Apache': 'Apache',
    'Window': 'Windows',
    'SSH': 'SSH',
    'Oracle': 'Oracle',
}

def get_module_name(name, predefined_keywords):
    # Iterate through the predefined keywords and check if they match the name
    for keyword in predefined_keywords:
        if keyword.lower() in name.lower():
            return keyword
    return ""  # Return empty string if no match is found

def append_data(doc, csv_row_no, data_to_append, predefined_keywords):
    # table - 0
    # finding_id
    finding_id_cell = doc.tables[csv_row_no * 8].rows[1].cells[0]
    finding_id_cell.text = data_to_append['finding_id']
    finding_id_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Prepend the required text to the description
    description = f"During Vulnerability assessment and Penetration testing we observed that, {data_to_append['description']}"
    doc.tables[csv_row_no * 8].rows[1].cells[1].text = description

    # table - 1
    # cvs_score
    doc.tables[csv_row_no * 8 + 1].rows[1].cells[0].text = data_to_append['cvs_score']

    # risk_rating (with color highlighting)
    cell = doc.tables[csv_row_no * 8 + 1].rows[1].cells[1]
    cell.text = data_to_append['risk_factor'].capitalize()

    risk_colors = {
        "critical": "800000",
        "high": "FF6600",
        "medium": "FFFF00",
        "low": "008000",
        "informational": "3366FF",
    }
    
    risk_color = risk_colors.get(data_to_append['risk_factor'].lower())
    if risk_color:
        tcPr = cell._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), risk_color)
        tcPr.append(tcShading)

    # remote_exploitability
    doc.tables[csv_row_no * 8 + 1].rows[1].cells[2].text = "Yes"

    # table - 2 (Affected Resource & Module Name)
    table = doc.tables[csv_row_no * 8 + 2]
    
    # Get affected resources as a list
    affected_resources = data_to_append["affected_resource"].split("\n")
    
    # Use new formatting function
    format_affected_resources(table, affected_resources)

    # Module Name - dynamically matched from predefined keywords
    module_name_cell = table.cell(1, 2)
    module_name_cell.text = get_module_name(data_to_append['name'], predefined_keywords)
    module_name_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # table - 3 (Security Risk)
    doc.tables[csv_row_no * 8 + 3].rows[1].cells[0].text = ""

    # table - 4 (Business Impact)
    doc.tables[csv_row_no * 8 + 4].rows[0].cells[1].text = ""

    # table - 5 (Workaround / Mitigation)
    solution = f"It is recommended: \n-To {data_to_append['mitigation']}"
    doc.tables[csv_row_no * 8 + 5].rows[1].cells[0].text = solution

    # table - 6 (Tool Used & References)
    doc.tables[csv_row_no * 8 + 6].rows[1].cells[0].text = "Nessus"
    doc.tables[csv_row_no * 8 + 6].rows[1].cells[2].text = data_to_append['references']

    # table - 7 (Proof of Concept)
    proof_of_concept_cell = doc.tables[csv_row_no * 8 + 7].rows[1].cells[0]
    existing_text = proof_of_concept_cell.text  # Preserve existing content
    proof_of_concept_cell.text = ""  # Clear the cell to format text properly

    # Adding "Figure 1 - Shows" with "Figure 1" in bold
    proof_paragraph = proof_of_concept_cell.paragraphs[0]
    run = proof_paragraph.add_run("\nFigure 1")
    run.bold = True
    proof_paragraph.add_run(" - Shows " + existing_text)  # Append previous content

def main():
    # Create a new document
    doc = Document()
    set_document_font(doc, "Helvetica", 10.5)
    grouped_vulnerabilities = {}  # Dictionary to store vulnerabilities and affected hosts

    # Add a title
    title = doc.add_heading('RNS DATA AUTOMATION', level=1)
    title.alignment = 1  # Center align
    csv_file_path = 'dataset.csv'

    # Counter for finding IDs
    finding_id_counter = 1

    # Open the CSV file and read its contents
    with open(csv_file_path, 'r', encoding='utf-8', errors='ignore') as csv_file:
        csv_reader = csv.reader(csv_file)
        next(csv_reader)  # Skip header row

        for csv_row_data in csv_reader:
            vulnerability_name = csv_row_data[7]  # Extract vulnerability name
            affected_host = f"{csv_row_data[4]}:{csv_row_data[6]}"  # Format as 'IP:Port'

            # If the vulnerability is already recorded, just append the affected host
            if vulnerability_name in grouped_vulnerabilities:
                grouped_vulnerabilities[vulnerability_name]["affected_resource"].add(affected_host)
            else:
                # Assign a new finding ID and store all required details
                grouped_vulnerabilities[vulnerability_name] = {
                    "name": vulnerability_name,
                    "finding_id": f"ABCXYZ-{finding_id_counter}",
                    "description": csv_row_data[9].split('.')[0],
                    "cvs_score": csv_row_data[2],
                    "risk_factor": csv_row_data[18],
                    "remote_exploitability": "Yes",
                    "affected_resource": {affected_host},  # Store as a set to avoid duplicates
                    "mitigation": csv_row_data[10],
                    "references": csv_row_data[11].split('\n')[0] if csv_row_data[11] else "",
                }
                finding_id_counter += 1  # Increment counter

    # Now, iterate through the grouped vulnerabilities and create tables
    # Now, iterate through the grouped vulnerabilities and create tables
    for index, (vulnerability_name, data_to_append) in enumerate(grouped_vulnerabilities.items()):
        # ✅ Fix: Strip spaces before joining affected resources
        data_to_append["affected_resource"] = "\n".join(res.strip() for res in sorted(data_to_append["affected_resource"]))

    # Create the table
        numbered_finding_name = f"{index + 1}. {vulnerability_name}"  
        create_table(doc, numbered_finding_name)
    
    # Add new page except for the last vulnerability
        if index < len(grouped_vulnerabilities) - 1:
            doc.add_section(WD_SECTION.NEW_PAGE)
        
        append_data(doc, index, data_to_append, KEYWORDS)

    # Save the document
    doc.save('output_document.docx')

if __name__ == "__main__":
    main()